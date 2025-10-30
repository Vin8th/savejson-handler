import json
from pathlib import Path

from logger import Logger

import docx
from docx.shared import Pt
from markdown2 import markdown
from bs4 import BeautifulSoup

from logger import Logger


class SaveJson:
    """
    This class saves the json response from the request to json file or word document.
    """

    sub_folder = None

    def __init__(self, log: Logger, base_dir: str):
        self.log = log
        self.base_dir = base_dir
        self.loc = None

    def save_json_result(self, data, filename, path, type):
        try:
            SaveJson.sub_folder = path

            base_dir = Path(self.base_dir)
            folder = (base_dir / path)
            folder.mkdir(parents=True, exist_ok=True)
            file_path = (folder / f"{filename}.{type}")
            n = result = data if (filename) == "type" else data
            result = data

            self.loc = file_path
            self.log.info(f"Saving file in : {file_path}")

            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=4)
                self.log.info("File saved")

        except Exception as e:
            self.log.error(f"Error saving: {e}")

    def json_markdown_to_word(self, data):
        md_content = self._json_to_markdown(data)

        html_content = markdown(md_content, extras=["tables", "fenced-code-blocks", "strike", "break-on-newline"])
        soup = BeautifulSoup(html_content, "html.parser")

        doc = docx.Document()

        doc.add_heading(f"{SaveJson.sub_folder} Metadata Documentation")

        for element in soup.descendants:
            if getattr(element, "name", None) is None:
                continue

            if element.name.startswith("h") and element.name[1].isdigit():
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
                continue

            if element.name in ["ul", "ol"]:
                style = "List Bullet" if element.name == "ul" else "List Number"
                for li in element.find_all("li", recursive=False):
                    doc.add_paragraph(li.get_text(), style=style)
                continue

            if element.name == "blockquote":
                doc.add_paragraph(element.get_text(), style="Intense Quote")
                continue

            if element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                cols = rows[0].find_all(["th", "td"])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["th", "td"])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
                doc.add_paragraph()
                continue

            if element.name == "pre":
                code_block = element.get_text().strip()
                p = doc.add_paragraph()
                run = p.add_run(code_block)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                continue

            p = doc.add_paragraph()
            for child in element.children:
                if isinstance(child, str):
                    p.add_run(child)
                    continue
                if child.name == "strong":
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == "em":
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == "code":
                    run = p.add_run(child.get_text())
                    run.font.name = "Courier New"
                else:
                    p.add_run(child.get_text() if hasattr(child, "get_text") else str(child))
            continue

        path = f"{self.base_dir}/{SaveJson.sub_folder}/{SaveJson.sub_folder}.docx"
        doc.save(path)
        print(f"Word document generated: {path}")

        return path

    def _json_to_markdown(self, json_string):
        """
        Converts a JSON string to a Markdown file.

        Args:
            json_string (str): A valid JSON string.
            output_path (str): Path to output markdown file, e.g., 'output.md'.
        """
        try:
            data = json.loads(json_string)
        except json.JSONDecodeError as e:
            return json_string

        def format_item(key, value, level=2):
            markdown = ""
            if isinstance(value, dict):
                markdown += f"\n{'#' * level} {key}\n"
                for k, v in value.items():
                    markdown += format_item(k, v, level + 1)
            elif isinstance(value, list):
                markdown += f"\n{'#' * level} {key}\n\n"
                for i, item in enumerate(value):
                    markdown += format_item(f"{key} item {i+1}", item, level + 1)
            else:
                markdown += f"\n{'#' * level} {key}\n\n{value}\n\n"
            return markdown

        md_content = "# JSON to Markdown\n\n"
        for k, v in data.items():
            md_content += format_item(k, v)

        return md_content
    
    def save_result_to_word(self, result, output_path):
        """
        Saves any Python result as plain text to a Word document.

        Args:
            result: The Python object or result to save (can be str, dict, list, etc.)
            output_path: The path for the output .docx file
        """
        doc = docx.Document()
        doc.add_heading('Python Result', level=1)
        doc.add_paragraph(str(result))
        doc.save(output_path)
        self.log.info(f"Word document generated: {output_path}")

# Usage:
# s = SaveJson(logger, base_dir)
# s.save_result_to_word(python_result, "output.docx")
